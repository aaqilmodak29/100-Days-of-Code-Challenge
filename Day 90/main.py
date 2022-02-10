import boto3
from docx import Document


polly_service = boto3.client('polly')


def word_to_text(file_name):
    document = Document(file_name)
    docum = []
    for i in range(0, len(document.paragraphs)):
        if len(document.paragraphs[i].text) > 0:
            if len(document.paragraphs[i].text) != 1 and document.paragraphs[i].text != '':
                docum.append(document.paragraphs[i].text)
    txt = str(docum).strip('[, ]')
    # print(txt)
    return txt

# def pdf_to_text(file_name):
#     pdf_read = PyPDF2.PdfFileReader(file_name)
#     for page_num in range(pdf_read.numPages):
#         page_obj = pdf_read.getPage(page_num)
#         txt = page_obj.extractText()
#         print(txt)
#
#         return txt


def play(text):
    response = polly_service.synthesize_speech(
        Text=text,
        VoiceId='Joanna',
        OutputFormat='mp3',
        Engine='neural'
    )

    body = response['AudioStream'].read()

    with open('output.mp3', mode='wb') as audio_file:
        audio_file.write(body)
        audio_file.close()


if __name__ == '__main__':
    file = input('Enter file name/path: ')
    play(word_to_text(file))

